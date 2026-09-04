"""
Generate an official Microsoft Word (.docx) deployment manual for eduManage360.
"""
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_manual():
    doc = docx.Document()

    # Configure Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Title Header ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("eduManage360 Platform")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)  # Deep Navy

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Standard Operating Procedure: Deploying on a New Laptop / Computer")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Metadata Box ─────────────────────────────────────────────────────────
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("System:", "eduManage360 Multi-School Management System (SaaS Edition)"),
        ("Target Operating Systems:", "Windows 10 / 11, macOS, Ubuntu Linux"),
        ("Python Version:", "Python 3.10, 3.11, or 3.12 (64-bit)"),
        ("Deployment Modes:", "Method A (Git Clone) & Method B (USB Flash Drive)")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = tbl.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = label
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9.5)
        c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 41, 59)
        c1.text = val
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "F8FAFC")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ── Section: Mandatory Prerequisite ──────────────────────────────────────
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Mandatory Prerequisite: Installing Python on the New Computer")
    r1.font.color.rgb = RGBColor(30, 58, 138)

    p_pre = doc.add_paragraph()
    p_pre.add_run("Before proceeding with either deployment method, you must install Python on the new computer:\n")
    p_pre.add_run("1. Download Python 3.10, 3.11, or 3.12 from ").font.size = Pt(10.5)
    p_pre.add_run("https://www.python.org/downloads/").font.bold = True
    p_pre.add_run("\n2. Run the installer.")

    # Warning Callout Box
    callout_tbl = doc.add_table(rows=1, cols=1)
    callout_cell = callout_tbl.rows[0].cells[0]
    set_cell_background(callout_cell, "FEF3C7")  # Amber highlight
    cp = callout_cell.paragraphs[0]
    c_run1 = cp.add_run("CRITICAL INSTALLATION REQUIREMENT (Windows):\n")
    c_run1.font.bold = True
    c_run1.font.color.rgb = RGBColor(180, 83, 9)
    c_run2 = cp.add_run("On the very first screen of the Python installer, check the box at the bottom:\n")
    c_run3 = cp.add_run("☑ Add python.exe to PATH\n")
    c_run3.font.bold = True
    c_run4 = cp.add_run("If this box is not checked, the 'python' and 'pip' commands will not work in your command prompt.")
    for r in [c_run2, c_run4]:
        r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── Section: Method A (8-Step Fresh Git Deployment) ──────────────────────
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Method A: The Complete 8-Step Fresh Git Deployment")
    r2.font.color.rgb = RGBColor(30, 58, 138)

    p_a_intro = doc.add_paragraph()
    p_a_intro.add_run("Use this method when you have internet access on the new laptop and wish to clone the latest codebase directly from GitHub.")

    steps_a = [
        ("Step 1: Install Git & Python", "Install Git from https://git-scm.com/ and Python 3.10+ (with 'Add to PATH' checked)."),
        ("Step 2: Clone the Project Repository", "Open PowerShell or Command Prompt and execute:\n   git clone https://github.com/Bismark-Duah/SMS.git\n   cd SMS"),
        ("Step 3: Create & Activate Virtual Environment", "Run in PowerShell:\n   python -m venv venv\n   .\\venv\\Scripts\\activate\n   (On macOS/Linux: source venv/bin/activate)"),
        ("Step 4: Install System Dependencies", "Install all backend packages:\n   pip install -r backend/requirements.txt"),
        ("Step 5: Create Environment File (.env)", "Create a file named backend/.env with the following lines:\n   ENVIRONMENT=offline_local\n   DATABASE_URL=sqlite:///./sql_app.db\n   SECRET_KEY=eduManage360_secure_key_2026\n   MNOTIFY_SENDER_ID=EDUMANAGE\n   SMS_PRIMARY_GATEWAY=mnotify"),
        ("Step 6: Initialize Database & Run Verification Harness", "Run the 51-suite verification harness to bootstrap database tables and default roles:\n   python verify_all.py\n   (Confirm that 'ALL 51 SYSTEM VERIFICATION TESTS PASSED!' is displayed)."),
        ("Step 7: Launch the Application Server", "Start the local server:\n   python run.py"),
        ("Step 8: Open in Browser & Log In", "Open Google Chrome, Brave, or Edge and navigate to:\n   http://localhost:8000/auth.html\n   Username: superadmin | Password: Superadmin123!")
    ]

    for title, desc in steps_a:
        p = doc.add_paragraph()
        run_t = p.add_run(f"• {title}: ")
        run_t.font.bold = True
        run_t.font.color.rgb = RGBColor(15, 23, 42)
        run_d = p.add_run(desc)
        run_d.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── Section: Method B (4-Step USB Flash Drive Transfer) ──────────────────
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. Method B: The 4-Step Direct USB / Flash Drive Transfer")
    r3.font.color.rgb = RGBColor(30, 58, 138)

    p_b_intro = doc.add_paragraph()
    p_b_intro.add_run("Use this method when you want to copy the entire system directly from your current laptop to the new laptop using a flash drive or external hard drive.")

    # Rule Note Box
    rule_tbl = doc.add_table(rows=1, cols=1)
    rule_cell = rule_tbl.rows[0].cells[0]
    set_cell_background(rule_cell, "EFF6FF")  # Light Blue
    rp = rule_cell.paragraphs[0]
    r_run1 = rp.add_run("GOLDEN RULE FOR USB TRANSFERS:\n")
    r_run1.font.bold = True
    r_run1.font.color.rgb = RGBColor(29, 78, 216)
    r_run2 = rp.add_run("Do NOT copy the 'venv/' folder (or delete it after copying). Virtual environments contain hardcoded file paths tied to your old laptop's username. Rebuilding it on the new laptop takes less than 60 seconds.")
    r_run2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    steps_b = [
        ("Step 1: Copy the 'SMS' Folder to Your Flash Drive", "Copy the entire SMS project folder from your current computer to your USB drive. You may exclude the 'venv' folder to save transfer time."),
        ("Step 2: Paste the Folder onto the New Laptop", "Paste the SMS folder into your desired destination on the new laptop (e.g. C:\\SMS or D:\\documents\\SMS)."),
        ("Step 3: Recreate the Python Virtual Environment", "Open PowerShell or Command Prompt inside the pasted SMS folder on the new laptop and run:\n   python -m venv venv\n   .\\venv\\Scripts\\activate\n   pip install -r backend\\requirements.txt"),
        ("Step 4: Verify Database & Start the System", "Run the verification harness and start the server:\n   python verify_all.py\n   python run.py\n   Open browser: http://localhost:8000/auth.html")
    ]

    for title, desc in steps_b:
        p = doc.add_paragraph()
        run_t = p.add_run(f"• {title}: ")
        run_t.font.bold = True
        run_t.font.color.rgb = RGBColor(15, 23, 42)
        run_d = p.add_run(desc)
        run_d.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── Section: 1-Click Master Cloud Mirroring ──────────────────────────────
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. Optional: 1-Click Master Cloud Mirroring from Render")
    r4.font.color.rgb = RGBColor(30, 58, 138)

    p_sync = doc.add_paragraph()
    p_sync.add_run("If your new laptop is deployed in a school office and you want to pull all registered schools, students, classes, and marks already stored on your live Render cloud database:\n\n"
                   "1. Open http://localhost:8000/super-admin.html on the new laptop.\n"
                   "2. Log in with: superadmin / Superadmin123!\n"
                   "3. Navigate to the 'System Operations & Tools' tab.\n"
                   "4. Click 'Open Cloud Synchronization Wizard'.\n"
                   "5. Enter Remote URL: https://sms-nald.onrender.com and password: Superadmin123!\n"
                   "6. Click 'Start Cloud Sync' — All live records will be downloaded and populated locally in under 5 seconds.")

    # ── Section: System Endpoints Reference Table ────────────────────────────
    h5 = doc.add_heading(level=1)
    r5 = h5.add_run("5. Quick Reference: System Portal URLs & Credentials")
    r5.font.color.rgb = RGBColor(30, 58, 138)

    tbl_ref = doc.add_table(rows=4, cols=3)
    tbl_ref.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Portal Description", "Local URL Endpoint", "Default Credentials"]
    for col_idx, h_text in enumerate(headers):
        cell = tbl_ref.rows[0].cells[col_idx]
        cell.text = h_text
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "0F172A")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    ref_data = [
        ("School Portal (Admin / Teachers / Bursar)", "http://localhost:8000/auth.html", "admin / admin123 (or school user)"),
        ("Super-Admin Executive Console", "http://localhost:8000/super-admin.html", "superadmin / Superadmin123!"),
        ("Interactive API Swagger Docs", "http://localhost:8000/docs", "FastAPI OpenAPI Specification")
    ]

    for row_idx, (p_name, url, creds) in enumerate(ref_data, start=1):
        r = tbl_ref.rows[row_idx]
        r.cells[0].text = p_name
        r.cells[1].text = url
        r.cells[2].text = creds
        for c in r.cells:
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(c, "F8FAFC" if row_idx % 2 == 1 else "FFFFFF")

    # Save Document
    out_path = os.path.abspath("eduManage360_Laptop_Deployment_Guide.docx")
    doc.save(out_path)
    print(f"Manual successfully generated at: {out_path}")

if __name__ == "__main__":
    create_manual()
